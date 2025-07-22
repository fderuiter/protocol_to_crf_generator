"""DOCX importer extracting text and tables."""

from __future__ import annotations

from io import StringIO
from pathlib import Path
import csv

from docx import Document


def load_docx(path: Path) -> tuple[str, list[str]]:
    """Load a .docx file and extract text and tables.

    Parameters
    ----------
    path:
        Path to the DOCX file.

    Returns
    -------
    tuple[str, list[str]]
        A tuple containing the document text and a list of tables as CSV strings.

    Raises
    ------
    ValueError
        If the supplied path does not end with ``.docx``.
    """
    if path.suffix.lower() != ".docx":
        raise ValueError("Unsupported file type")

    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    tables: list[str] = []
    for table in doc.tables:
        buffer = StringIO()
        writer = csv.writer(buffer)
        for row in table.rows:
            writer.writerow([cell.text.replace("\n", " ").strip() for cell in row.cells])
        tables.append(buffer.getvalue().replace("\r\n", "\n").strip())
        buffer.close()

    return text, tables
