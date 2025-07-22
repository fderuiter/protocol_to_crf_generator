from pathlib import Path

from protocol_to_crf_generator.ingestion import load_docx

from docx import Document


def _create_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Protocol Title")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header1"
    table.rows[0].cells[1].text = "Header2"
    table.rows[1].cells[0].text = "Value1"
    table.rows[1].cells[1].text = "Value2"
    doc.save(path)


def test_load_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    _create_sample_docx(file_path)
    text, tables = load_docx(file_path)
    assert "Protocol Title" in text
    assert tables == ["Header1,Header2\nValue1,Value2"]


def test_invalid_extension(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("not a docx")
    try:
        load_docx(file_path)
    except ValueError:
        pass
    else:
        assert False, "Expected ValueError"
